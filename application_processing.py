import random
import struct
from enum import Enum

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import Table

import json

steel_hardness = {
    "Сталь 20": (137, 153),
    "Сталь 10": (125, 143),
    "Сталь 25Л": (121, 151),
    "Сталь 20Л": (111, 156),  # Обратите внимание: значения для двух вхождений "Сталь 20Л" различаются. В словаре останется последнее (111, 156)
    "Сталь 3сп": (111, 156),
    "Сталь 09Г2С": (135, 152),
    "Сталь 15": (137, 148),
    "Сталь 25": (162, 170),
    "Сталь 30ХМ": (200, 250)
}

class SectionType(Enum):
    TUBE = "труба"
    SHURF = "Шурф"
    ZMS = "ЗМС"

class Section:
    def __init__(self, number, type : SectionType, picket, du, area_nominal, steel, thick):
        self.number = number
        self.type = type # TODO сейчас шурф будет без номера
        self.picket = picket
        self.du = du # диаметр
        self.area_nominal = area_nominal

        if steel not in steel_hardness.keys():
            raise Exception(f"bad steel type: {steel}")
        self.steel = steel

        if not ((self.area_nominal - 0.5) <= thick <= self.area_nominal): # TODO нужно ли равно
            raise Exception(f"bad thick: {thick} | Must be in range ({self.area_nominal - 0.5}; {self.area_nominal})")
        self.thick = thick

        self.environment = None

        self.diam_measure_results = []
        self.min_diam = None
        self.thick_measure_results = []
        self.min_thick = None

# priloj 5: 2st - 35-... по порядку.
#           4-5st - константы


    def set_values(self):
        if self.type == SectionType.ZMS:
            count_measures = 4
            count_measures_rows = 1
        else:
            count_measures_rows = 3
            count_measures = 6

        for i in range(count_measures):
            self.diam_measure_results.append(random.randrange(*steel_hardness[self.steel]))

        self.min_diam = min(self.diam_measure_results)

        for i in range(count_measures_rows):
            self.thick_measure_results.append([])
            for j in range(count_measures):
                self.thick_measure_results[i].append(random.randrange(int(self.thick) * 10, int(self.area_nominal) * 10) / 10.)

        self.min_thick = min([min(_) for _ in self.thick_measure_results])



def set_cell_format(cell, default_paragraph):
    cell.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = default_paragraph.alignment
        paragraph.style = default_paragraph.style
        for run in paragraph.runs:
            run.font.size = default_paragraph.runs[0].font.size
            run.font.name = default_paragraph.runs[0].font.name
            run.font.color.rgb = default_paragraph.runs[0].font.color.rgb
            run.style = default_paragraph.runs[0].style
            run.italic = default_paragraph.runs[0].italic
            run.underline = default_paragraph.runs[0].underline
            run.bold = default_paragraph.runs[0].bold


import csv
import json
from collections import defaultdict

# json для отбраковочного значения
def parse_csv_to_json(csv_file, json_file):
    # Чтение всех строк CSV
    with open(csv_file, 'r', encoding='utf-8') as f:
        reader = csv.reader(f, delimiter=';')
        rows = list(reader)

    # Структура для накопления данных
    data = defaultdict(lambda: defaultdict(lambda: defaultdict(dict)))

    i = 0
    while i < len(rows):
        row = rows[i]
        # Поиск начала блока: строка, содержащая "сталь" в первой ячейке или второй
        # В первой ячейке часто пусто, а во второй - заголовок
        if len(row) > 1 and 'сталь' in row[1].lower():
            # Извлекаем материал, давление и среду из ячеек 3,4,5 (индексы 3,4,5)
            # Пример: row[3] = "Сталь 20", row[4] = "4", row[5] = "НГЖС"
            material = row[3].strip()  # "20"
            pressure = row[4].strip().replace(',', '.')      # "4.0"
            medium_raw = row[5].strip()
            # Приводим среду к единому виду: "вода" для "вода"/"инертная жидкость"
            if medium_raw.lower() in ['вода', 'инертная жидкость']:
                medium = 'вода'
            elif medium_raw.lower() == 'нгжс':
                medium = 'НГЖС'
            elif medium_raw.lower() == 'газ':
                medium = 'газ'
            else:
                medium = medium_raw  # на случай других вариантов

            # Теперь нужно найти строки с диаметрами и результатами внутри этого блока
            # Пропускаем строки до "Наружный диаметр элемента, мм"
            j = i + 1
            while j < len(rows) and 'Наружный диаметр элемента' not in rows[j][1]:
                j += 1
            if j >= len(rows):
                break
            diam_row = rows[j]
            # Диаметры начинаются с индекса 1 (первый столбец - описание)
            diameters = diam_row[2:]  # список значений диаметров, чередуются труба/отвод

            # Далее ищем строку с "- рассчитанное"
            calc_row = None
            accept_row = None
            k = j + 1
            while k < len(rows):
                if rows[k] and rows[k][1] and '- рассчитанное' in rows[k][1]:
                    calc_row = rows[k]
                if rows[k] and rows[k][1] and '- принятое' in rows[k][1]:
                    accept_row = rows[k]
                    break
                k += 1

            if calc_row is None or accept_row is None:
                # Если не нашли, пропускаем блок
                i = k
                continue

            calc_values = calc_row[2:]  # рассчитанные значения
            accept_values = accept_row[2:]  # принятые значения

            # Заполняем структуру
            # diameters, calc_values, accept_values имеют одинаковую длину
            # Каждая пара столбцов (чётный и нечётный) соответствуют одному диаметру
            num_pairs = len(diameters) // 2
            for pair in range(num_pairs):
                # Индексы для трубы (2*pair) и отвода (2*pair + 1)
                idx_tr = 2 * pair
                idx_ot = 2 * pair + 1

                # Диаметр для этой пары (труба и отвод имеют одинаковый диаметр)
                diam = diameters[idx_tr].strip()
                # Преобразуем в строку, чтобы использовать как ключ JSON
                diam_key = diam

                # Значения для трубы
                try:
                    calc_tr = float(calc_values[idx_tr].replace(',', '.'))
                except:
                    calc_tr = None
                try:
                    acc_tr = float(accept_values[idx_tr].replace(',', '.'))
                except:
                    acc_tr = None

                # Значения для отвода
                try:
                    calc_ot = float(calc_values[idx_ot].replace(',', '.'))
                except:
                    calc_ot = None
                try:
                    acc_ot = float(accept_values[idx_ot].replace(',', '.'))
                except:
                    acc_ot = None

                # Сохраняем
                data[material][medium][pressure][diam_key] = {
                    'труба': {
                        'рассчитанное': calc_tr,
                        'принятое': acc_tr
                    },
                    'отвод': {
                        'рассчитанное': calc_ot,
                        'принятое': acc_ot
                    }
                }

            # Перемещаем указатель после обработанного блока
            i = k + 1
        else:
            i += 1

    # Преобразуем defaultdict в обычный dict и записываем JSON
    result = json.loads(json.dumps(data))
    with open(json_file, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

def add_row_pril_12_2(sections : list[Section], table : Table):
    row_start_index = 3
    for obj_number in range(len(sections)):
        if not (sections[obj_number].type == SectionType.ZMS):
            row_index = row_start_index + obj_number * 3
            for _ in range(3):
                table.add_row()
            for coll_index in range(5):
                table.row_cells(row_index)[coll_index].merge(table.row_cells(row_index + 2)[coll_index])
            count_rows = 3
        else:
            row_index = row_start_index + obj_number
            table.add_row()
            count_rows = 1

        curr_section = sections[obj_number]
        table.row_cells(row_index)[0].text = f'{curr_section.number}\n{curr_section.type.value}'
        table.row_cells(row_index)[1].text = f'{curr_section.picket}'

        table.row_cells(row_index)[2].text = f'{curr_section.du}'
        table.row_cells(row_index)[3].text = f"{str(float(curr_section.area_nominal)).replace('.', ',')}"
        # table.row_cells(row_index)[4].text = 'ПОТОМ'

        with open('data/otbrak_table.json', 'br') as f:
            otbrak_data = json.load(f)

        if curr_section.type != SectionType.ZMS:
            # TODO сейчас давление, диаметр заданы строго, надо как-то заменить

            # TODO также проверить как оно с дробными значениями, где они нужны и тд.
            # Также некоторые числа строки в json, мб поменять

            # TODO шурф == отвод???
            if curr_section.type.value == 'Шурф':
                table.row_cells(row_index)[4].text = str(otbrak_data[curr_section.steel]['вода']['16'][str(curr_section.du)]['отвод']['принятое']).replace('.', ',')
            else:
                table.row_cells(row_index)[4].text = str(otbrak_data[curr_section.steel]['вода']['16'][str(curr_section.du)][curr_section.type.value]['принятое']).replace('.', ',')
        else:
            table.row_cells(row_index)[4].text = 'HZ'
            # TODO сделать для ЗМС, нужно спросить


        for i in range(count_rows):
            for j in range(len(curr_section.thick_measure_results[i])):
                table.row_cells(row_index + i)[5 + j].text = str(curr_section.thick_measure_results[i][j])

        # setting style

        default_paragraph = table.row_cells(0)[0].paragraphs[0]

        for i in range(count_rows):
            for j in range(len(table.row_cells(row_index + i))):
                set_cell_format(table.row_cells(row_index + i)[j], default_paragraph)


def add_row_pril_13(sections: list[Section], table: Table):
    row_start_index = 3
    for obj_number in range(len(sections)):
        row_index = row_start_index + obj_number
        table.add_row()

        curr_section = sections[obj_number]
        table.row_cells(row_index)[0].text = f'{curr_section.number}'
        table.row_cells(row_index)[1].text = f'{curr_section.type.value}'
        table.row_cells(row_index)[2].text = f'{curr_section.picket}'

        table.row_cells(row_index)[3].text = f'{curr_section.du}'
        table.row_cells(row_index)[4].text = f"{curr_section.steel}"
        # table.row_cells(row_index)[4].text = 'ПОТОМ'

        with open('data/otbrak_table.json', 'br') as f:
            otbrak_data = json.load(f)

        for j in range(len(curr_section.diam_measure_results)):
            table.row_cells(row_index)[5 + j].text = str(curr_section.diam_measure_results[j])

        # setting style

        default_paragraph = table.row_cells(0)[0].paragraphs[0]

        for j in range(len(table.row_cells(row_index))):
            set_cell_format(table.row_cells(row_index)[j], default_paragraph)
