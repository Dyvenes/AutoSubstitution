from docx import Document
from docx.shared import RGBColor
import re
from typing import Dict, Any, List
import os

# app.py
from flask import Flask, render_template, request, send_file, session
from datetime import datetime
import os
from docx import Document
import io
import secrets

app = Flask(__name__)
app.secret_key = secrets.token_hex(16)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Создаем папки если их нет
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs('static', exist_ok=True)


class WordTemplateProcessor:
    def __init__(self, template_path: str):
        self.doc = Document(template_path)
        self.replacements = {}

    def set_replacements(self, data: Dict[str, Any]):
        """Установка данных для замены"""
        self.replacements = data

    def merge_runs_with_placeholders(self, paragraph):
        """
        Объединяет Runs, если они содержат части плейсхолдеров
        Например: [{'{'}, {'{компания'}, {'}}'}] -> ['{{компания}}']
        """
        runs = list(paragraph.runs)
        i = 0

        while i < len(runs) - 1:
            current_text = runs[i].text
            next_text = runs[i + 1].text

            # Проверяем, являются ли эти Runs частями плейсхолдера
            # Ищем паттерны типа: { + {компания, {{ + компания, компания + }}, { + компания + }
            combined = current_text + next_text

            # Проверяем, является ли combined или его части плейсхолдером
            is_placeholder_part = (
                    '{{' in combined or '}}' in combined or
                    any(ph in combined for ph in self.replacements.keys()) or
                    (current_text.endswith('{') and next_text.startswith('{')) or
                    (current_text.endswith('}') and next_text.startswith('}'))
            )

            if is_placeholder_part:
                # Объединяем Runs
                runs[i].text = combined
                # Удаляем следующий Run
                runs[i + 1].text = ''
                # Удаляем пустой Run из списка
                runs.pop(i + 1)
            else:
                i += 1

        # Очищаем пустые Runs
        for run in paragraph.runs[:]:
            if run.text == '':
                p_element = run._element
                p_element.getparent().remove(p_element)

    def smart_replace_in_paragraph(self, paragraph):
        """
        Умная замена с предварительным объединением Runs
        """
        # Шаг 1: Объединяем Runs, которые могут быть частями плейсхолдеров
        self.merge_runs_with_placeholders(paragraph)

        # Шаг 2: Выполняем замену
        original_runs = list(paragraph.runs)
        if not original_runs:
            return

        # Проверяем, есть ли вообще плейсхолдеры в параграфе
        paragraph_text = paragraph.text
        if not any(key in paragraph_text for key in self.replacements.keys()):
            return

        # Шаг 3: Выполняем замену с сохранением форматирования
        paragraph.clear()

        for run in original_runs:
            new_text = run.text
            # Заменяем все плейсхолдеры в этом Run
            for key, value in self.replacements.items():
                if key in new_text:
                    new_text = new_text.replace(key, str(value))

            if new_text:  # Добавляем только непустой текст
                new_run = paragraph.add_run(new_text)
                # Копируем все свойства форматирования
                self.copy_run_formatting(run, new_run)

    def copy_run_formatting(self, source_run, target_run):
        """Копирует форматирование из одного Run в другой"""
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size

        # Копируем цвет, если он установлен
        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

        # Дополнительные свойства
        if hasattr(source_run.font, 'highlight_color'):
            target_run.font.highlight_color = source_run.font.highlight_color

    def process_document(self):
        """Обработка всего документа"""
        # Обрабатываем обычные параграфы
        for paragraph in self.doc.paragraphs:
            self.smart_replace_in_paragraph(paragraph)

        # Обрабатываем таблицы
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.smart_replace_in_paragraph(paragraph)

        # Обрабатываем другие элементы
        self.process_headers_footers()

    def process_headers_footers(self):
        """Обработка верхних и нижних колонтитулов"""
        for section in self.doc.sections:
            # Верхние колонтитулы
            for paragraph in section.header.paragraphs:
                self.smart_replace_in_paragraph(paragraph)
            # Нижние колонтитулы
            for paragraph in section.footer.paragraphs:
                self.smart_replace_in_paragraph(paragraph)

    def save(self, output_path: str):
        """Сохранение документа"""
        self.doc.save(output_path)
        print(f"✓ Документ сохранен: {output_path}")


def save_template_to_server(file):
    """Сохранение загруженного шаблона"""
    filename = f"template_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    return filepath


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        try:
            # Получаем данные из формы
            form_data = {
                'company_name': request.form.get('company_name', ''),
                'contract_number': request.form.get('contract_number', ''),
                'director_name': request.form.get('director_name', ''),
                'address': request.form.get('address', ''),
                'contract_date': request.form.get('contract_date', ''),
                'contract_type': request.form.get('contract_type', '')
            }

            # Сохраняем в сессии для повторного использования
            session['form_data'] = form_data

            # Обрабатываем загруженный файл
            if 'template_file' not in request.files:
                return render_template('index.html',
                                       error='Пожалуйста, загрузите файл шаблона',
                                       form_data=form_data)

            template_file = request.files['template_file']
            if template_file.filename == '':
                return render_template('index.html',
                                       error='Файл не выбран',
                                       form_data=form_data)

            if not template_file.filename.endswith('.docx'):
                return render_template('index.html',
                                       error='Поддерживаются только файлы .docx',
                                       form_data=form_data)

            # Сохраняем шаблон
            template_path = save_template_to_server(template_file)

            # Создаем словарь для замены
            replacements = {
                '{{filial}}': form_data['company_name'],
                '{{company}}': form_data['contract_number'],
                '{{DIRECTOR_NAME}}': form_data['director_name'],
                '{{ADDRESS}}': form_data['address'],
                '{{CONTRACT_DATE}}': form_data['contract_date'],
                '{{CONTRACT_TYPE}}': form_data['contract_type']
            }

            # Обрабатываем шаблон
            processor = WordTemplateProcessor(template_path)

            processor.set_replacements(replacements)

            processor.process_document()

            # Сохраняем результат в памяти
            output = io.BytesIO()
            processor.save("test_file_name.docx")
            output.seek(0)

            # Генерируем имя файла
            output_filename = f"Договор_{form_data['company_name']}_{datetime.now().strftime('%Y%m%d')}.docx"

            # Очищаем сессию
            session.pop('form_data', None)

            # Отправляем файл пользователю
            return send_file(output,
                             as_attachment=True,
                             download_name=output_filename,
                             mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        except Exception as e:
            return render_template('index.html',
                                   error=f'Ошибка: {str(e)}',
                                   form_data=form_data if 'form_data' in locals() else {})

    # GET запрос - показываем форму
    form_data = session.get('form_data', {
        'company_name': '',
        'contract_number': '',
        'director_name': '',
        'address': '',
        'contract_date': datetime.now().strftime('%Y-%m-%d'),
        'contract_type': 'поставки'
    })

    return render_template('index.html', form_data=form_data)


@app.route('/download_sample')
def download_sample():
    """Скачивание примерного шаблона"""
    # Создаем примерный шаблон
    doc = Document()

    # Заголовок
    doc.add_heading('ДОГОВОР {{CONTRACT_TYPE}} №{{CONTRACT_NUMBER}}', 0)

    # Дата
    doc.add_paragraph(f'Дата: {{CONTRACT_DATE}}')

    # Стороны договора
    doc.add_heading('1. СТОРОНЫ ДОГОВОРА', 1)
    doc.add_paragraph(f'Исполнитель: {{COMPANY_NAME}}')
    doc.add_paragraph(f'Адрес: {{ADDRESS}}')
    doc.add_paragraph(f'Генеральный директор: {{DIRECTOR_NAME}}')

    # Предмет договора
    doc.add_heading('2. ПРЕДМЕТ ДОГОВОРА', 1)
    doc.add_paragraph('Исполнитель обязуется выполнить работы, а Заказчик принять и оплатить их.')

    # Стоимость и порядок расчетов
    doc.add_heading('3. СТОИМОСТЬ И ПОРЯДОК РАСЧЕТОВ', 1)

    # Добавляем таблицу
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Заголовки таблицы
    headers = table.rows[0].cells
    headers[0].text = 'Наименование работ'
    headers[1].text = 'Количество'
    headers[2].text = 'Стоимость'

    # Данные таблицы
    data_rows = table.rows[1].cells
    data_rows[0].text = '{{CONTRACT_TYPE}}'
    data_rows[1].text = '1'
    data_rows[2].text = '100 000 руб.'

    # Подписи
    doc.add_paragraph('\n\n')
    doc.add_paragraph('Исполнитель: _________________ {{DIRECTOR_NAME}}')
    doc.add_paragraph('\n')
    doc.add_paragraph('Заказчик: _________________')

    # Сохраняем в BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(output,
                     as_attachment=True,
                     download_name='Пример_шаблона.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')



# Пример использования с таблицами
def example_complex_replacement():
    template_path = "sto-01-538-2017.docx"

    processor = WordTemplateProcessor(template_path)

    # Данные для замены
    replacements = {
        '{{филиал}}': 'НАЗВАНИЕ ФИЛИАЛА',
        '{{company}}': 'НАЗВАНИЕ КОМПАНИИ',
        '{{kind_of_number}}': 'СЕРИЙНИК',
        '{{какой-то старый номер}}': 'СТАРЫЙ СЕРИЙНИК',
    }

    processor.set_replacements(replacements)
    processor.process_document()
    processor.save("готовый_договор.docx")

    return template_path, "готовый_договор.docx"

#
# # Основной скрипт
# if __name__ == "__main__":
#
#     app.run(debug=True, port=5000)
#     # print("=== Обработчик Word-документов ===")
#     #
#     # # Пример 1: Обработка одного файла
#     # print("\n1. Обработка одного файла:")
#     # example_complex_replacement()
#     #
#     # # Пример 2: Пакетная обработка (раскомментируйте для использования)
#     # # print("\n2. Пакетная обработка:")
#     # # batch_processing_example()
#     #
#     # # Пример 3: Из командной строки
#     """
#     Использование:
#     python word_processor.py --template шаблон.docx --output результат.docx \
#         --replace '{"name": "Иван", "date": "2024-01-15"}'
#     """