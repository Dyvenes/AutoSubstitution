# main.py
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from typing import Optional
import os
import io
import uuid
import shutil
from datetime import datetime
from pathlib import Path
from docx import Document

# Создаем приложение FastAPI
app = FastAPI(
    title="Генератор Word-документов",
    description="Веб-приложение для заполнения шаблонов Word",
    version="1.0.0"
)

# Настройка CORS для доступа с любого домена
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # В продакшене укажите конкретные домены
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "output"
STATIC_DIR = BASE_DIR / "static"

TEMPLATE_PATH = "sto-01-538-2017.docx"

UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)
STATIC_DIR.mkdir(exist_ok=True)

# Подключаем статические файлы и шаблоны
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

class WordTemplateProcessor:
    def __init__(self, template_path: str):
        self.doc = Document(template_path)
        self.replacements = {}

    def set_replacements(self, data: dict):
        """Установка данных для замены"""
        self.replacements = data

    def merge_runs_with_placeholders(self, paragraph):
        """Объединяет Runs, если они содержат части плейсхолдеров"""
        runs = list(paragraph.runs)
        i = 0

        while i < len(runs) - 1:
            current_text = runs[i].text
            next_text = runs[i + 1].text
            combined = current_text + next_text

            is_placeholder_part = (
                    '{{' in combined or '}}' in combined or
                    any(ph in combined for ph in self.replacements.keys()) or
                    (current_text.endswith('{') and next_text.startswith('{')) or
                    (current_text.endswith('}') and next_text.startswith('}'))
            )

            if is_placeholder_part:
                runs[i].text = combined
                runs[i + 1].text = ''
                runs.pop(i + 1)
            else:
                i += 1

        # Очищаем пустые Runs
        for run in paragraph.runs[:]:
            if run.text == '':
                p_element = run._element
                p_element.getparent().remove(p_element)

    def smart_replace_in_paragraph(self, paragraph):
        """Умная замена с предварительным объединением Runs"""
        self.merge_runs_with_placeholders(paragraph)
        original_runs = list(paragraph.runs)

        if not original_runs:
            return

        paragraph_text = paragraph.text
        if not any(key in paragraph_text for key in self.replacements.keys()):
            return

        paragraph.clear()

        for run in original_runs:
            new_text = run.text
            for key, value in self.replacements.items():
                if key in new_text:
                    new_text = new_text.replace(key, str(value))

            if new_text:
                new_run = paragraph.add_run(new_text)
                self.copy_run_formatting(run, new_run)

    def copy_run_formatting(self, source_run, target_run):
        """Копирует форматирование из одного Run в другой"""
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size

        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    def process_document(self):
        """Обработка всего документа"""
        for paragraph in self.doc.paragraphs:
            self.smart_replace_in_paragraph(paragraph)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.smart_replace_in_paragraph(paragraph)

    def get_bytes(self):
        """Возвращает документ в виде bytes"""
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output.getvalue()

@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    print("IN GET")
    return templates.TemplateResponse("index.html", {
        "request": request,
        "today": datetime.now().strftime("%Y-%m-%d")
    })


@app.post("/generate")
async def generate_document(
        company_name: str = Form(...),
        contract_number: str = Form(...),
        director_name: str = Form(...),
        address: str = Form(...),
        contract_date: str = Form(...),
        contract_type: str = Form(...),
        template_file: UploadFile = File(...)
):
    """Генерация документа с заменой плейсхолдеров"""

    try:
        # Создаем словарь для замены
        replacements = {
            '{{filial}}': company_name,
            '{{company}}': contract_number,
            '{{kind_of_number}}': director_name,
            '{{old_kind_of_number}}': address,
            '{{CONTRACT_DATE}}': contract_date,
            '{{CONTRACT_TYPE}}': contract_type
        }

        # Обрабатываем документ
        try:
            processor = WordTemplateProcessor(TEMPLATE_PATH)
            processor.set_replacements(replacements)
            processor.process_document()

            # Генерируем имя файла
            output_filename = f"Договор_{company_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = OUTPUT_DIR / output_filename

            # Сохраняем результат
            processor.doc.save(output_path)

            # Отправляем файл пользователю
            return FileResponse(
                path=output_path,
                filename=output_filename,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Ошибка обработки документа: {str(e)}"
            )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка сервера: {str(e)}"
        )

@app.get("/download_sample")
async def download_sample():

    sample_path = "sto-01-538-2017.docx"

    return FileResponse(
        path=sample_path,
        filename="Пример_шаблона.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/check_network")
async def check_network(request: Request):
    """Проверка доступности сервера"""
    client_host = request.client.host if request.client else "unknown"
    return {
        "status": "server is running",
        "client_ip": client_host,
        "server_ip": "178.157.138.159",
        "time": datetime.now().isoformat()
    }


# main.py - обновленный блок запуска
if __name__ == "__main__":
    import uvicorn
    import socket


    # Получаем локальный IP
    def get_local_ip():
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            ip = s.getsockname()[0]
            s.close()
            return ip
        except:
            return "127.0.0.1"


    local_ip = get_local_ip()
    print("=" * 50)
    print(f"Сервер запущен!")
    print(f"Локальный доступ:  http://localhost:8000")
    print(f"Сетевой доступ:    http://{local_ip}:8000")
    print(f"Ваш внешний IP:    http://178.157.138.159:8000")
    print("=" * 50)

    # Запуск с настройками для внешнего доступа
    uvicorn.run(
        "main:app",
        host="0.0.0.0",  # Принимать соединения со всех интерфейсов
        port=8000,
        reload=True,
        # Добавляем для лучшей сетевой совместимости
        access_log=True,
        log_level="info"
    )