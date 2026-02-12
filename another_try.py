from docx.oxml import parse_xml
from flask import Flask, request, send_file, render_template, jsonify

from pathlib import Path
from datetime import datetime
import io

from docx import Document

import pandas as pd

from db import DatabaseManager

import logging
import sys

# Настройка логирования (добавьте после импортов)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),  # В systemd
    ]
)
logger = logging.getLogger(__name__)

app = Flask(
    __name__,
    static_folder="static",
    template_folder="templates"
)

logger.info("initializing db")
db = DatabaseManager()
logger.info("db is ready")

BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "output"
STATIC_DIR = BASE_DIR / "static"
TEMPLATE_PATH = BASE_DIR / "templates/template_file.docx"
DATA_DIR = BASE_DIR / "data"

UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)
STATIC_DIR.mkdir(exist_ok=True)

months = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
}


class WordTemplateProcessor:
    def __init__(self, template_path: str):
        self.doc = Document(template_path)
        self.replacements = {}

    def set_replacements(self, data: dict):
        self.replacements = data

    def process_headers_footers(self):
        """Обработка всех колонтитулов в документе"""
        # Обработка верхних колонтитулов (headers)
        for section in self.doc.sections:
            # Header первой страницы
            if section.first_page_header is not None:
                for paragraph in section.first_page_header.paragraphs:
                    self.smart_replace_in_paragraph(paragraph)
                # Обработка таблиц в колонтитулах
                for table in section.first_page_header.tables:
                    self.process_table_in_container(table)

            # Обычный верхний колонтитул
            for paragraph in section.header.paragraphs:
                self.smart_replace_in_paragraph(paragraph)
            for table in section.header.tables:
                self.process_table_in_container(table)

            # Нижние колонтитулы (footers)
            if section.first_page_footer is not None:
                for paragraph in section.first_page_footer.paragraphs:
                    self.smart_replace_in_paragraph(paragraph)
                for table in section.first_page_footer.tables:
                    self.process_table_in_container(table)

            for paragraph in section.footer.paragraphs:
                self.smart_replace_in_paragraph(paragraph)
            for table in section.footer.tables:
                self.process_table_in_container(table)

    def process_table_in_container(self, table):
        """Обработка таблиц в колонтитулах"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.smart_replace_in_paragraph(paragraph)

    def merge_runs_with_placeholders(self, paragraph):
        runs = list(paragraph.runs)
        i = 0
        while i < len(runs) - 1:
            current_text = runs[i].text
            next_text = runs[i + 1].text
            combined = current_text + next_text
            is_placeholder_part = (
                    '{{' in combined or '}}' in combined or
                    any(ph in combined for ph in self.replacements.keys()) or
                    (current_text.endswith('{') and next_text.startswith('{')) or
                    (current_text.endswith('}') and next_text.startswith('}'))
            )
            if is_placeholder_part:
                runs[i].text = combined
                runs[i + 1].text = ''
                runs.pop(i + 1)
            else:
                i += 1

        # Очищаем пустые Runs
        for run in paragraph.runs[:]:
            if run.text == '':
                p_element = run._element
                p_element.getparent().remove(p_element)

    def smart_replace_in_paragraph(self, paragraph):
        if not paragraph.runs:
            return

        paragraph_text = paragraph.text
        if not any(key in paragraph_text for key in self.replacements.keys()):
            return

        runs = list(paragraph.runs)
        # logger.info([i.text for i in runs])

        #logger.info("STARTING", ['' for i in paragraph.runs])
        if not any('{{' in i.text for i in runs):
            return

        # logger.info("DO REPLACE")

        for i in range(len(runs)):
            if '{{' in runs[i].text and '}}' in runs[i].text:
                # logger.info("FOUND ONE VAR IN:", runs[i].text)
                for key, val in self.replacements.items():
                    if key in runs[i].text:
                        # logger.info("VAR IS:", key)
                        paragraph.runs[i].text = paragraph.runs[i].text.replace(key, str(val))
                        # logger.info("REPLACE SUCCESS")

        i = 0
        while i < len(runs) - 2:
            if not runs[i]:
                continue

            new_text = runs[i].text
            if '{{' == new_text.strip():
                """
                TODO сейчас не надежно ищет, основываясь на том, что всегда 
                делит отдельными ранами. Надо отслеживать, если встретились два {{ подряд, то рэйсить ошибку
                
                когда удаляется }} могут теряться пробелы 
                """
                index_to_write = i

                i += 1

                new_text = runs[i].text


                # paragraph.runs[i].text = ''

                var = str()

                while '}}' != new_text.strip():
                    var += new_text
                    paragraph.runs[i].text = ''
                    i += 1
                    new_text = runs[i].text
                paragraph.runs[i].text = paragraph.runs[i].text.replace('}}', '')

                for key, value in self.replacements.items():
                    clear_key = key.replace('{{', '')
                    clear_key = clear_key.replace('}}', '')
                    if clear_key == var:
                        paragraph.runs[index_to_write].text = str(value)
                        break
            i += 1


    def copy_run_formatting(self, source_run, target_run):
        """Копирует форматирование из одного Run в другой"""
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size

        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    def process_document(self):
        """Обработка всего документа"""
        for paragraph in self.doc.paragraphs:
            self.smart_replace_in_paragraph(paragraph)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.smart_replace_in_paragraph(paragraph)

        self.process_headers_footers()

    def replace_table_by_index(self, table_index, new_table_xml):
        tables = self.doc.tables

        if table_index < len(tables):
            # Получаем элемент таблицы для замены
            old_table_element = tables[table_index]._tbl

            # Получаем родительский элемент
            parent = old_table_element.getparent()

            # Создаем новую таблицу
            if isinstance(new_table_xml, str):
                new_table_element = parse_xml(new_table_xml)
            else:
                # Если передан Table объект
                new_table_element = new_table_xml._tbl

            # Заменяем старую таблицу новой
            parent.replace(old_table_element, new_table_element)

            print(f"Таблица {table_index} успешно заменена")
        else:
            print(f"Таблица с индексом {table_index} не найдена")

    def get_bytes(self):
        """Возвращает документ в виде bytes"""
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output.getvalue()


@app.route("/", methods=["GET"])
def read_root():
    logger.info("IN GET")
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate_document():
    """Генерация документа с заменой плейсхолдеров"""
    try:
        logger.info("GENERATING")
        filename = request.form.get("filename")  # TODO потом генерировать автоматически??

        grafic = request.files.get("graf_file")
        report_number = request.form.get("TO_number")

        pipline_type = request.form.get("pipline_type")
        pipline_name = request.form.get("pipline_name")

        temperature = request.form.get("temperature")
        pressure_work = request.form.get("pressure_work")
        pressure_project = request.form.get("pressure_project")
        insulation = request.form.get('insulation')
        logger.info("INSULATION: " + str(insulation))
        anticor = request.form.get("anticor")
        inside_cover = request.form.get("inside_cover")
        welding = request.form.get("welding")
        project_documentation = request.form.get("project_documentation")
        installation_company = request.form.get("installation_company")
        logger.info("GOT VALUES")

        # template_file вы сейчас игнорируете и всегда берете TEMPLATE_PATH
        # Если нужно использовать загруженный шаблон, можно сохранить его и
        # подставлять путь вместо TEMPLATE_PATH.

        curr_date = datetime.now().strftime("%d.%m.%Y")
        str_curr_date = f"{datetime.now().day} {months[datetime.now().month]} {datetime.now().year} года"
        curr_year = str(datetime.now().year)
        year_short = curr_year[-2:]

        logger.info("CURR_YEAR: " + str(curr_year))
        # --------------------------------------------------------
        # EXCEL

        import time

        start = time.time()

        # GRAFIC TABLE
        df = pd.read_excel(grafic, nrows=100, dtype=str, engine='openpyxl')

        csv = df.to_csv(sep=';').split('\r\n')
        csv = [i.split(';') for i in csv]

        row_index = 0

        logger.info("GOT REPORT_NUMBER: " + str(report_number))
        for row in range(len(csv)):
            if report_number in csv[row]:
                row_index = row

        end = time.time()

        logger.info("Time = " + str(end - start))

        # excel = win32com.client.Dispatch("Excel.Application")
        # excel.Visible = False  # Скрыть Excel
        #
        # wb = excel.Workbooks.Open(r'worl_files/grafic.xlsx')
        # ws = wb.Worksheets(1)  # Первый лист

        deposit = csv[row_index][3]
        workshop = csv[row_index][4]
        inventory_number = csv[row_index][5]
        # pipline_name = csv[row_index][7]
        length_of_pipline = csv[row_index][11]
        length_of_area = csv[row_index][12]
        wall_diam = csv[row_index][9]
        wall_thic = str(float(csv[row_index][10])).replace('.', ',')
        wall_params = f"{wall_diam}x{wall_thic}"
        year_of_commissioning = datetime.fromisoformat(csv[row_index][13]).year
        year_of_using = datetime.now().year - year_of_commissioning
        diagnostic_date = datetime.fromisoformat(csv[row_index][19]).strftime("%d.%m.%Y")

        leader_surname = csv[row_index][21]

        # HZ NOMERNAYA TABLE
        #number_table = request.files.get("report_table_file")

        #enter_lay = pd.read_excel(number_table, sheet_name="Ввод данных", nrows=20, dtype=str, engine='openpyxl')
        #logger.info("FILE READED")
        #enter_lay_csv = enter_lay.to_csv(sep=';').split('\r\n', )
        #logger.info(enter_lay_csv)
        #enter_lay_csv = [i.split(';') for i in enter_lay_csv]

        #full_pipline_name = enter_lay_csv[0][2]
        # ----------------------------------------------------------
        # DATABASE

        logger.info("LEADER SURNAME: " + str(leader_surname))

        leader = None # Employer class
        employees = db.get_all_employees()
        for employer in employees:
            logger.info("db surname: " + str(employer.surname))
            print(employer.surname.strip() == leader_surname.strip())
            if employer.surname.strip() == leader_surname.strip():
                leader = employer
                break
        else:
            logger.info("Leader NOT found")
            raise "Leader not found"

        logger.info("Leader found")
        team_number = leader.team_number
        leader_full = leader.surname + " " + leader.name + " " + leader.lastname
        leader_short = leader.name[0] + ". " + leader.lastname[0] + ". " + leader.surname
        leader_position = leader.position
        leader_license = leader.license

        # team_members = [] если вдруг команда будет состоять больше чем из 2 человек
        worker = None # Employer class
        for employer in employees:
            if employer.team_number == team_number and employer.id != leader.id:
                worker = employer
                break

        worker_full = worker.surname + " " + worker.name + " " + worker.lastname
        worker_short = worker.name[0] + ". " + worker.lastname[0] + ". " + worker.surname
        worker_position = worker.position
        worker_license = worker.license
        logger.info("ALL DATA GOT")
        instrument_table = leader.instrument_table
        logger.info("L_F: " + str(leader_full))
        logger.info("L_S: " + str(leader_short))
        logger.info("L_Pos: " + str(leader_position))
        logger.info("L_Lic: " + str(leader_license))

        logger.info("W_F: " + str(worker_full))
        logger.info("W_S: " + str(worker_short))
        logger.info("W_Pos: " + str(worker_position))
        logger.info("W_Lic: " + str(worker_license))

        # ----------------------------------------------------------
        replacements = {
            '{{curr_year}}': curr_year,
            '{{report_number}}': report_number,
            '{{rep_num}}': report_number,
            '{{year_short}}': year_short,

            '{{pipline_name}}': pipline_name,  # временно
            '{{pipline_type}}': pipline_type,
            '{{full_pipline_name}}': f'{pipline_type} «{pipline_name}»',

            '{{inventory_number}}': inventory_number,
            '{{deposit}}': deposit,
            '{{workshop}}': workshop,
            '{{diagnostic_date}}': diagnostic_date,
            '{{curr_date}}': curr_date,
            '{{str_curr_date}}': str_curr_date,  # день сделать двойным числом всегда
            '{{length_of_area}}': length_of_area,
            '{{length_of_pipline}}': length_of_pipline,
            '{{wall_params}}': wall_params,
            '{{wall_diam}}': wall_diam,
            '{{wall_thic}}': wall_thic,
            '{{year_of_commissioning}}': year_of_commissioning, # эксплуатации
            '{{years_of_using}}': year_of_using,

            '{{leader_full}}': leader_full, # Иванов Иван Иванович
            '{{leader_short}}': leader_short, # Иванов И. И.
            '{{leader_position}}': leader_position,
            '{{leader_license}}': leader_license,

            '{{worker_full}}': worker_full,
            '{{worker_short}}': worker_short,
            '{{worker_position}}': worker_position,
            '{{worker_license}}': worker_license,

            '{{temperature}}': temperature,
            '{{pressure_work}}': pressure_work,
            '{{pressure_project}}': pressure_project,
            '{{insulation}}': insulation,
            '{{anticor}}': anticor,
            '{{inside_cover}}': inside_cover,
            '{{welding}}': welding,
            '{{project_documentation}}': project_documentation,
            '{{installation_company}}': installation_company
        }

        processor = WordTemplateProcessor(str(TEMPLATE_PATH))
        logger.info("CREATED FILE")
        processor.replace_table_by_index(3, instrument_table)
        logger.info("TABLE REPLACED")
        processor.set_replacements(replacements)
        logger.info("SET REPLACEMENTS")
        processor.process_document()
        logger.info("PROCESS")

        output_filename = f'{filename}.docx'
        logger.info("FILE NAME: " + str(output_filename))

        output_path = OUTPUT_DIR / output_filename
        processor.doc.save(output_path)

        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return {"detail": f"Ошибка сервера: {str(e)}"}, 500



@app.route("/process_graf_file", methods=["POST"])
def process_graf_file():
    logger.info("PROCESSING FILE")
    file = request.files.get("graf_file")

    file.save("data/graf.xlsx") # ?если файл не подходящего формата, то надо его удалить?
    # создает файл, даже если его по сути нет
    with open("data/date_manager.txt", "w+") as manager: # может лучше сделать csv
        manager.write(f"graf.xlsx;{file.filename}\n") # _TODO сейчас вроде перезаписывается весь файл


    df = pd.read_excel(file, nrows=100, dtype=str, engine='openpyxl')

    csv = df.to_csv(sep=';').split('\r\n')
    csv = [i.split(';') for i in csv]

    row = 0
    col = 0
    for i in range(len(csv)):
        if 'ТО №' in csv[i]:
            row = i + 1 # в текущем - заголовок
            col = csv[i].index('ТО №')
            break
    else:
        return {
            "success": False,
            "error": "Неизвестный формат файла"
        }, 500

    logger.info("ROW: " + str(row) + " COL: " + str(col))

    options = [csv[i_r][col] for i_r in range(row, len(csv)) if (len(csv[i_r]) > 2 and csv[i_r][col] != '')]
    logger.info("NUMERS: " + str(options))
    return {
        "success": True,
        "data": options,
    }

@app.route("/preload_files", methods=["GET"])
def preload_files():
    with open("data/date_manager.txt", 'r') as manager:
        files_names = manager.read().split('\n')
        for i in range(len(files_names)):
            if files_names[i] == '':
                files_names.pop(i)
        if len(files_names) == 0 or not any(line != '' for line in files_names):
            return {
                "success": True,
                "files": None
            }
        files = []
        for file_name in files_names:
            files.append(file_name.split(';')[0])
        logger.info("PRELOADING FILENAMES: " + str(files))
        return {
            "success": True,
            "files": files
        }

@app.route("/upload_file", methods=["GET"])
def upload_file():

    filename = request.args.get('name')
    logger.info(str(filename))

    file_path = DATA_DIR / filename
    logger.info("UPLOADING FILE: " + str(file_path))

    original_filename = 'file not exist'
    with open('data/date_manager.txt', 'r') as manager:
        file_lines = manager.read().split('\n')
        for line in file_lines:
            if filename in line:
                original_filename = line.split(';')[1]

    return send_file(
        file_path,
        as_attachment=True,
        download_name=original_filename
    )

@app.route("/get_teams_list", methods=["GET"])
def get_teams_list():
    db.get_all_employees()


@app.route("/download_sample", methods=["GET"])
def download_sample():
    sample_path = TEMPLATE_PATH
    return send_file(
        sample_path,
        as_attachment=True,
        download_name="template.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.route("/check_network", methods=["GET"])
def check_network():
    """Проверка доступности сервера"""
    client_host = request.remote_addr or "unknown"
    return {
        "status": "server is running",
        "client_ip": client_host,
        "server_ip": "178.157.138.159",
        "time": datetime.now().isoformat()
    }


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
